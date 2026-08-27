"""Render a TR-doc JSON into a styled .docx matching the golden format,
and a parallel Markdown file for quick review."""
from __future__ import annotations
import re
from pathlib import Path

import docx as docxlib
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

BREAKER = "---------------------------------------"

# A list item that already carries its own marker: "1.", "2)", "3 -", "•", "- ".
_ENUMERATED = re.compile(r"^\s*\d+\s*[.)\-:]\s+")


def _list_item(doc, text: str, style: str = "List Bullet"):
    """Add one list line, with EXACTLY ONE marker in front of it.

    The agenda is required to be numbered "1."…"N." so it mirrors the numbered Key
    Takeaways (guardrails enforces it), and the model numbers the recap to match. Both
    were then written with style="List Bullet", which puts a bullet glyph in front of
    the number — so the document read "• 1. Buffering", two markers deep, in the recap,
    the agenda and the Key Takeaways alike.

    Fixing this in the prompt would be fighting the guardrail that demands the numbers.
    So the RENDERER decides: text that already numbers itself keeps its own number and
    loses the glyph (an indented paragraph, no bullet); everything else is bulleted as
    before. The model's numbers are used verbatim rather than Word's automatic ones,
    which continue counting across consecutive lists and would number the agenda 6..10
    straight after a five-line recap.
    """
    body = str(text)
    if _ENUMERATED.match(body):
        p = doc.add_paragraph(body.strip(), style="List Paragraph")
        return p
    return doc.add_paragraph(body, style=style)


def _md_list(items) -> str:
    """One Markdown line per item, with EXACTLY ONE marker — the .docx rule (see
    _list_item) applied to the review pane and the .md preview, so what a reviewer
    approves on screen is what the document shows."""
    out = []
    for i in items or []:
        body = str(i)
        out.append(body.strip() if _ENUMERATED.match(body) else f"- {body}")
    return "\n".join(out)


def _labelled(doc, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label} ")
    run.bold = True
    p.add_run(str(value))
    return p


def code_lines(block: dict) -> list[str]:
    """A code block's lines, as rendered. One helper so the writer, the page estimate and
    the guardrails all agree about what "a line of code" is — the page ceiling is a hard
    gate and it is measured against this."""
    return str(block.get("code") or "").replace("\r\n", "\n").split("\n")


def walkthrough_text(block: dict) -> list[str]:
    """The prose of a code block — what the learner reads about the snippet.

    This, not the code, is the code slide's writing: it is what the prose/bullet mix gate
    counts, what the banned-phrase checks read, and what the recording estimate treats as
    spoken. `for (int i = 0; i < n; i++)` is not a sentence and must not be graded as one.
    """
    out = []
    for w in block.get("walkthrough") or []:
        if isinstance(w, dict) and str(w.get("text") or "").strip():
            out.append(str(w["text"]).strip())
    return out


def _walk_label(w: dict) -> str:
    lines = str(w.get("lines") or "").strip()
    return f"Line {lines}" if lines and "-" not in lines else f"Lines {lines}"


def _render_content(doc, blocks):
    for block in blocks or []:
        t = block.get("type")
        if t == "text":
            doc.add_paragraph(block.get("text", ""))
        elif t == "bullets":
            for item in block.get("items", []):
                _list_item(doc, item)
        elif t == "code":
            # Monospace, ONE PARAGRAPH PER LINE with the space-after removed, so the
            # snippet reads as a block rather than as double-spaced prose. The walkthrough
            # follows as a list, because it is what the learner reads about the code.
            for line in code_lines(block):
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.left_indent = Pt(18)
                run = para.add_run(line or " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            for w in block.get("walkthrough") or []:
                if not isinstance(w, dict) or not str(w.get("text") or "").strip():
                    continue
                _list_item(doc, f"{_walk_label(w)} — {str(w['text']).strip()}")
        elif t == "table":
            cols = block.get("columns", [])
            rows = block.get("rows", [])
            if not cols:
                continue
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            for i, c in enumerate(cols):
                cell = table.rows[0].cells[i]
                cell.text = str(c)
                for para in cell.paragraphs:
                    for r in para.runs:
                        r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for i in range(len(cols)):
                    cells[i].text = str(row[i]) if i < len(row) else ""


def write_docx(doc_json: dict, out_path: Path) -> Path:
    d = docxlib.Document()
    n = doc_json.get("session_no", "")
    title = doc_json.get("session_title", "")

    d.add_heading(f"Session {n} : {title}", level=1)
    d.add_paragraph()

    recap = doc_json.get("recap")
    if recap:
        d.add_heading(
            f"RECAP: Session {recap['prev_session_no']} : {recap['prev_session_name']}",
            level=2)
        for b in recap.get("bullets", []):
            _list_item(d, b)

    d.add_heading("Agenda for Today's Session", level=2)
    for a in doc_json.get("agenda", []):
        _list_item(d, a)

    for sec in doc_json.get("sections", []):
        d.add_paragraph()
        d.add_heading(f"{BREAKER} SECTION {sec['index']}: {sec['name']} {BREAKER}", level=2)
        for s in sec.get("slides", []):
            d.add_heading(f"Slide {s['n']}: {s['title']}", level=3)
            _labelled(d, "Heading:", s.get("heading", ""))
            _labelled(d, "Subheading:", s.get("subheading", ""))
            p = d.add_paragraph()
            p.add_run("Content:").bold = True
            _render_content(d, s.get("content", []))
            if s.get("analogy"):
                _labelled(d, "Analogy:", s["analogy"])
            if s.get("visual_guidance"):
                _labelled(d, "Visual Guidance:", s["visual_guidance"])
            if s.get("speaker_notes"):
                _labelled(d, "Speaker Notes:", f"\"{s['speaker_notes']}\"")
            d.add_paragraph()

    d.add_heading("Key Takeaways", level=2)
    for k in doc_json.get("key_takeaways", []):
        _list_item(d, k)

    d.add_paragraph()
    up = doc_json.get("upcoming_session")
    if up:
        _labelled(d, "Upcoming Session :", up)

    closing = d.add_paragraph(doc_json.get("closing", "Thank You  |  All the Best"))
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in closing.runs:
        r.bold = True
        r.font.size = Pt(14)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(out_path))
    return out_path


def _content_blocks(content) -> list[str]:
    """Markdown blocks for one slide's ordered content (text / bullets / table)."""
    out = []
    for block in content or []:
        bt = block.get("type")
        if bt == "text":
            out.append(block.get("text", ""))
        elif bt == "bullets":
            out.append(_md_list(block.get("items", [])))
        elif bt == "code":
            lang = str(block.get("language") or "").strip()
            out.append(f"```{lang}\n" + "\n".join(code_lines(block)) + "\n```")
            walk = [f"- {_walk_label(w)} — {str(w['text']).strip()}"
                    for w in (block.get("walkthrough") or [])
                    if isinstance(w, dict) and str(w.get("text") or "").strip()]
            if walk:
                out.append("\n".join(walk))
        elif bt == "table":
            cols = block.get("columns", [])
            tbl = ["| " + " | ".join(cols) + " |",
                   "| " + " | ".join(["---"] * len(cols)) + " |"]
            tbl += ["| " + " | ".join(str(c) for c in row) + " |"
                    for row in block.get("rows", [])]
            out.append("\n".join(tbl))
    return out


def _slide_blocks(s: dict) -> list[str]:
    """Markdown blocks for one slide, matching the .docx layout."""
    blocks = [f"### Slide {s['n']}: {s['title']}",
              f"**Heading:** {s.get('heading','')}  \n"
              f"**Subheading:** {s.get('subheading','')}",
              "**Content:**"]
    blocks += _content_blocks(s.get("content", []))
    # Each extra is its OWN block (blank-line separated) so Analogy / Visual
    # Guidance / Speaker Notes are clearly distinguishable in the preview.
    if s.get("analogy"):
        blocks.append(f"**Analogy:** {s['analogy']}")
    if s.get("visual_guidance"):
        blocks.append(f"**Visual Guidance:** {s['visual_guidance']}")
    if s.get("speaker_notes"):
        blocks.append(f"**Speaker Notes:** \"{s['speaker_notes']}\"")
    return blocks


def chunk_to_markdown(kind: str, fragment: dict) -> str:
    """Render ONE guided chunk (opening or section) to Markdown for the review pane."""
    blocks: list[str] = []
    if kind == "opening":
        recap = fragment.get("recap")
        if recap:
            blocks.append(f"## RECAP: Session {recap.get('prev_session_no','')} : "
                          f"{recap.get('prev_session_name','')}")
            if recap.get("bullets"):
                blocks.append(_md_list(recap["bullets"]))
        else:
            blocks.append("_(First session — no recap.)_")
        blocks.append("## Agenda for Today's Session")
        if fragment.get("agenda"):
            blocks.append(_md_list(fragment["agenda"]))
    else:  # section
        sec = fragment.get("section", fragment)
        idx = sec.get("index", "")
        blocks.append(f"## {BREAKER} SECTION {idx}: {sec.get('name','')} {BREAKER}")
        for s in sec.get("slides", []):
            blocks += _slide_blocks(s)
    return "\n\n".join(blocks) + "\n"


def write_markdown(doc_json: dict, out_path: Path) -> Path:
    # Each element of `blocks` is one Markdown block. Blocks are joined with a
    # BLANK line ("\n\n") so headings, paragraphs, lists and tables each render
    # as separate elements. (A single newline is NOT a line break in Markdown —
    # that was collapsing the whole doc into one flowing paragraph in the UI.)
    # Tight label pairs use a trailing "  " (hard break) to stack without a gap.
    blocks: list[str] = []
    n, title = doc_json.get("session_no", ""), doc_json.get("session_title", "")
    blocks.append(f"# Session {n} : {title}")

    recap = doc_json.get("recap")
    if recap:
        blocks.append(f"## RECAP: Session {recap['prev_session_no']} : {recap['prev_session_name']}")
        if recap.get("bullets"):
            blocks.append(_md_list(recap["bullets"]))

    blocks.append("## Agenda for Today's Session")
    if doc_json.get("agenda"):
        blocks.append(_md_list(doc_json["agenda"]))

    for sec in doc_json.get("sections", []):
        blocks.append(f"## {BREAKER} SECTION {sec['index']}: {sec['name']} {BREAKER}")
        for s in sec.get("slides", []):
            blocks += _slide_blocks(s)

    blocks.append("## Key Takeaways")
    if doc_json.get("key_takeaways"):
        blocks.append(_md_list(doc_json["key_takeaways"]))
    if doc_json.get("upcoming_session"):
        blocks.append(f"**Upcoming Session :** {doc_json['upcoming_session']}")
    blocks.append(f"**{doc_json.get('closing','Thank You  |  All the Best')}**")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n\n".join(blocks) + "\n", encoding="utf-8")
    return out_path
